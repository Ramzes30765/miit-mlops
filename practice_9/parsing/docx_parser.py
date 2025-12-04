from __future__ import annotations

from pathlib import Path
from typing import List
from uuid import uuid4

import docx

from .base import BaseParser, ParsedDocument
from .text_normalizer import normalize_text


class DocxParser(BaseParser):
    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def parse(self, path: Path | str) -> ParsedDocument:
        path = Path(path)

        document = docx.Document(str(path))

        paragraphs: List[str] = []

        for p in document.paragraphs:
            text = "".join(run.text for run in p.runs)
            text = text.strip()
            if not text:
                continue
            paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)
        full_text = normalize_text(full_text)

        return ParsedDocument(
            doc_id=str(uuid4()),
            text=full_text,
            metadata={
                "source_path": str(path),
                "file_name": path.name,
                "file_type": "docx",
                "num_paragraphs": len(paragraphs),
            },
        )
